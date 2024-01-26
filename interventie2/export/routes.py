import os
from flask import current_app, Blueprint, render_template, redirect, url_for, send_file, request
from flask_login import login_user, logout_user, current_user, login_required
from interventie2.models import User, Worksession, QuestionSet, Instrument, Option, Answer, Selection, Question, Process
from interventie2.models import db, generate_secret_key
from interventie2.forms import ExportWorksessionForm
from interventie2.classes import Advisor
from docx import Document
from docx.shared import Cm
import simplejson as json
from interventie2.export.export_modules import add_title, add_worksession_info, add_remarks, add_worksession_process, add_suggestions_table, add_answers, add_instrument, add_calculation
from interventie2.classes import Advisor


export = Blueprint('export', __name__,
                 template_folder='templates',
                 static_folder='static',
                 static_url_path='assets')



@export.route('/')
@login_required
def index():
    instruments = Instrument.query.order_by(Instrument.name)
    question_sets = QuestionSet.query.order_by(QuestionSet.name)
    worksessions = Worksession.query.order_by(Worksession.name)
    return render_template('export/index.html', instruments=instruments, question_sets=question_sets, worksessions=worksessions)

# These functions export things to Word. These functions call separate export modules from another file.
# This is in general unelegantly programmed. The called functions append requested things to a document file.
# The document file is an empty Word file with required styles. This file is copied and sent to the user.
# This is just a terrible way to do this, and I apologize.

@export.route('/worksession/<int:worksession_id>', methods=['GET', 'POST'])
@login_required
def worksession(worksession_id):
    worksession = Worksession.query.get(worksession_id)
    form = ExportWorksessionForm()
    advisor = Advisor(worksession=worksession, instruments=Instrument.query.all())

    print(current_app.static_folder)
    if form.validate_on_submit():
        result = Document(os.path.join(current_app.static_folder, 'export', 'Template.docx'))
        output_file = os.path.join(current_app.static_folder, 'export', 'Temp.docx')

        add_title(result, worksession.name)
        result.add_page_break()

        add_worksession_info(result, worksession)
        if form.export_technical_info.data:
            add_worksession_process(result, worksession)
        if len(form.remarks.data) > 0:
            add_remarks(result, form.remarks.data)
        result.add_page_break()

        add_answers(result, worksession)
        result.add_page_break()

        add_suggestions_table(result, advisor.get_sorted_instruments())
        result.add_page_break()

        result.add_heading("Selectie catalogus", level=1)
        for instrument, score in advisor.get_sorted_instruments():
            if form.export_all_instruments.data or (score > advisor.get_highest_score() - worksession.mark_top_instruments):
                add_instrument(result, instrument)
                if form.export_calculations.data:
                    add_calculation(result, advisor.explain_score(instrument))

        result.save(output_file)
        return send_file(output_file, as_attachment=True, download_name=f'{worksession.name}.docx')

    elif request.method == 'GET':
        form.export_tags.data = worksession.show_tags
        form.export_all_instruments.data = worksession.show_rest_instruments
    return render_template('export/worksession.html', worksession=worksession, form=form)


@export.route('/instrument/<int:instrument_id>')
@login_required
def instrument(instrument_id):
    """Export instrument to Word."""
    instrument = Instrument.query.get(instrument_id)
    result = Document(os.path.join(current_app.static_folder, 'export', 'Template.docx'))
    output_file = os.path.join(current_app.static_folder, 'export', 'Temp.docx')   

    result.add_heading('Instrumentbeschrijving', level=1)
    add_instrument(result, instrument)
    
    result.save(output_file)
    return send_file(output_file, as_attachment=True, download_name=f'{instrument.name}.docx')


@export.route('/catalog')
@login_required
def catalog():
    result = Document(os.path.join(current_app.static_folder, 'export', 'Template.docx'))
    output_file = os.path.join(current_app.static_folder, 'export', 'Temp.docx')   

    add_title(result, 'Catalogus')

    for instrument in Instrument.query.order_by(Instrument.name):
        add_instrument(result, instrument, start_level=1)
    
    result.save(output_file)
    return send_file(output_file, as_attachment=True, download_name='catalogus.docx')


# These functions export to JSON. 

@export.route('/instrument_to_json/<int:instrument_id>')
@login_required
def instrument_to_json(instrument_id):
    instrument = Instrument.query.get(instrument_id)
    output_file = os.path.join(current_app.static_folder, 'export', 'Temp.json')   

    instrument_json = {
        "filetype": "instrument",
        "filetype_version": current_app.config["FILETYPE_VERSION"],
        "name": instrument.name,
        "introduction": instrument.introduction,
        "description": instrument. description,
        "referenced_elsewhere": instrument.referenced_elsewhere,
        "reference_link": instrument.reference_link,
        "considerations": instrument.considerations,
        "examples": instrument.examples,
        "links": instrument.links,
    }
    tags = []
    for tag in instrument.tags:
        tag_json = {
            "name": tag.tag.name,
            "weight": tag.weight,
            "multiplier": tag.multiplier}
        tags.append(tag_json)
    instrument_json['tags'] = tags

    with open(output_file, 'w') as output:
        output.write(json.dumps(instrument_json, use_decimal=True, indent=2))
    return send_file(output_file, as_attachment=True, download_name=f'{instrument.name}.json')


@export.route('/question_set_to_json/<int:question_set_id>')
@login_required
def question_set_to_json(question_set_id):
    question_set = QuestionSet.query.get(question_set_id)
    output_file = os.path.join(current_app.static_folder, 'export', 'Temp.json')  

    # General info
    question_set_json = {
        "filetype": "question_set",
        "filetype_version": current_app.config["FILETYPE_VERSION"],
        "name": question_set.name,
        "default_process": question_set.default_process_id,
        "description": question_set.description
    }

    # Forbidden and mandatory tags
    forbidden_tags = []
    for tag in question_set.forbidden_tags:
        tag_json = {
            "name": tag.name,
        }
        forbidden_tags.append(tag_json)
    question_set_json['forbidden_tags'] = forbidden_tags
    mandatory_tags = []
    for tag in question_set.mandatory_tags:
        tag_json = {
            "name": tag.name,
        }
        mandatory_tags.append(tag_json)
    question_set_json['mandatory_tags'] = mandatory_tags

    # Questions and characteristics
    questions = []
    for question in sorted(question_set.questions, key=lambda order: getattr(order, 'order')):
        question_json = {
            "name": question.name,
            "description": question.description,
            "allow_motivation": question.allow_motivation,
            "allow_choice": question.allow_choice,
            "allow_multiselect": question.allow_multiselect,
            "allow_weight": question.allow_weight
        }
        required_active_tags = []
        for tag in question.required_active_tags:
            required_active_tags_json = {
                "name": tag.name,
            }
            required_active_tags.append(required_active_tags_json)
        options = []
        for option in sorted(question.options, key=lambda order: getattr(order, 'order')):
            option_json = {
                "name": option.name
            }
            tags_json = []
            for tag in option.tags:
                tags_json.append(tag.name)
            option_json['tags'] = tags_json
            options.append(option_json)
        question_json['required_active_tags'] = required_active_tags
        question_json['options'] = options
        questions.append(question_json)
    question_set_json['questions'] = questions

    with open(output_file, 'w') as output:
        output.write(json.dumps(question_set_json, use_decimal=True, indent=2))
    return send_file(output_file, as_attachment=True, download_name=f'{question_set.name}.json')
