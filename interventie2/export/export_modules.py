from docx.shared import Cm

def add_title(document, title):
    document.add_heading(title, level=0)


def add_remarks(document, remarks):
    document.add_heading('Opmerkingen bij deze export', level=1)
    document.add_paragraph(remarks)


def add_worksession_info(document, worksession):
    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = f'Verslag {worksession.name} op {worksession.date.strftime("%d %B %Y")}.'

    document.add_heading(worksession.name, level=1)
    document.add_paragraph(f'Projectnummer {worksession.project_number}')

    document.add_heading('Beschrijving', level=2)
    document.add_paragraph(worksession.description)
    document.add_heading('Beoogd doel', level=2)
    document.add_paragraph(worksession.effect)

    # if worksession.link_to_page is not None:
    ## Commented out: who needs a link in a Word document?
    #     document.add_heading('Link naar project', level=2)
    #     document.add_paragraph(worksession.link_to_page)

    document.add_heading('Informatie over de sessie', level=1)

    document.add_heading('Datum', level=2)
    document.add_paragraph(worksession.date.strftime("%d %B %Y"))
    document.add_heading('Deelnemers', level=2)
    document.add_paragraph(worksession.participants)


def add_worksession_process(document, worksession):
    document.add_heading('Werkwijze', level=2)
    document.add_paragraph(f'De gebruikte tool was {worksession.question_set.name}. De tool werd gebruikt met het volgende proces: {worksession.process.name}.')


def add_answers(document, worksession):
    document.add_heading(f'Antwoorden op {worksession.question_set.name}', level=1)
    for question in sorted(worksession.question_set.questions, key=lambda order: getattr(order, 'order')):
        if question.is_category:
            document.add_heading(question.name, level=2)
            if question.description != "":
                document.add_paragraph(question.description)
        else:
            document.add_paragraph(question.name, style='Vraag')
            if question.description != "":
                document.add_paragraph(question.description)
            for option in question.options:
                if worksession.is_option_selected(option):
                    document.add_paragraph(option.name, style='Optie')
            for answer in worksession.answers:
                if answer.question == question:
                    document.add_paragraph(answer.motivation)


def add_suggestions_table(document, suggestions):
    document.add_heading('Advies op basis van de antwoorden', level=1)
    document.add_paragraph('Deze lijst van suggesties voor interventies is gebaseerd op de antwoorden gekozen in de sessie. In het volgende hoofdstuk staan de instrumenten verder uitgewerkt.')
    
    table = document.add_table(rows=1, cols=2, style='interventie.afm.nl')
    
    header = table.rows[0].cells
    header[0].text = 'Instrument'
    header[1].text = 'Score'
    header[1].width = Cm(3.0)
    for instrument, score in suggestions:
        row = table.add_row().cells
        row[0].text = instrument.name
        row[1].text = str(round(score,1))
        

def add_instrument(document, instrument, start_level=2):
    document.add_heading(instrument.name, level=start_level)
    document.add_paragraph(instrument.introduction)

    table = document.add_table(rows=1, cols=2, style='interventie.afm.nl')
    
    header = table.rows[0].cells

    row = table.add_row().cells
    row[0].text = 'Beschrijving'
    row[1].text = instrument.description
    
    row = table.add_row().cells
    row[0].text = 'Overwegingen bij gebruik'
    row[1].text = instrument.considerations
    
    row = table.add_row().cells
    row[0].text = 'Voorbeelden'
    row[1].text = instrument.examples

    row = table.add_row().cells
    row[0].text = 'Tags'
    tags = row[1].add_paragraph()
    for tag_assignment in instrument.tags:
        if tag_assignment.multiplier >= 1:
            tags.add_run(f'{tag_assignment.tag.name}: {tag_assignment.multiplier}', style='PlusTag')
        else:
            tags.add_run(f'{tag_assignment.tag.name}: {tag_assignment.multiplier}', style='MinTag')
        tags.add_run(' ')

   
def add_calculation(document, explanation):
    document.add_heading('Berekening', level=3)
    # if explanation.get 'forbidden_tags_found'] is not None:
    #     pass
    # if explanation['mandatory_tags_not_found'] is not None:
    #     pass
    # for tag in explanation.tags:
    #     pass

def add_intervention_plans(document, worksession):
    document.add_heading('Interventieplannen', level=1)
    for intervention_plan in worksession.plan:
        document.add_heading(intervention_plan.description, level=2)
        document.add_paragraph(f'Interventieplan gemaakt op {intervention_plan.date.strftime("%d %B %Y")}.')
        document.add_paragraph(intervention_plan.conclusion)


        table = document.add_table(rows=1, cols=2, style='interventie.afm.nl')
        header = table.rows[0].cells
        header[0].text = 'Instrument'
        header[1].text = 'Omschrijving'
        for instrument_choice in intervention_plan.instruments:
            row = table.add_row().cells
            row[0].text = instrument_choice.instrument.name
            row[1].text = instrument_choice.instrument.description
